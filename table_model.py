from pathlib import Path
import random
import colorsys

import pandas as pd
from PySide2.QtCore import QAbstractTableModel, Qt
from PySide2.QtGui import QColor, QIcon
from PySide2 import QtCore, QtWidgets

import civiltools_rc
# import matplotlib.cm as cm
# from matplotlib.colors import Normalize
# import matplotlib
import FreeCADGui as Gui

civiltools_path = Path(__file__).absolute().parent

low = 'cyan'
intermediate = 'yellow'
high = 'red'

# def color_map_color(value, norm, cmap_name='rainbow'):
#     cmap = cm.get_cmap(cmap_name)  # PiYG
#     rgb = cmap(norm(abs(value)))[:3]  # will return rgba, we take only first 3 so we get rgb
#     color = matplotlib.colors.rgb2hex(rgb)
#     return color

class ResultsModel(QAbstractTableModel):
    '''
    MetaClass Model for showing Results
    '''
    def __init__(self, data, headers):
        QAbstractTableModel.__init__(self)
        self.df = pd.DataFrame(data, columns=headers)
        
    def rowCount(self, parent=None):
        return self.df.shape[0]

    def columnCount(self, parent=None):
        return self.df.shape[1]

    def data(self, index, role=Qt.DisplayRole):
        return None

    def headerData(self, col, orientation, role):
        if role != Qt.DisplayRole:
            return
        if orientation == Qt.Horizontal:
            return self.headers[col]
        return int(col + 1)

    # def sort(self, col, order):
    #     """Sort table by given column number."""
    #     self.layoutAboutToBeChanged.emit()
    #     self.df.sort_values(
    #         by=self.df.columns[col],
    #         ascending = order == Qt.AscendingOrder,
    #         kind="mergesort",
    #         inplace=True,
    #     )
    #     self.df.reset_index(drop=True, inplace=True)
    #     self.layoutChanged.emit()

class DriftModel(ResultsModel):
    def __init__(self, data, headers):
        super(DriftModel, self).__init__(data, headers)
        self.df = self.df[[
            'Story',
            'OutputCase',
            'Max Drift',
            'Avg Drift',
            'Allowable Drift'
        ]]
        self.df = self.df.astype({'Max Drift': float, 'Avg Drift': float, 'Allowable Drift': float})
        self.headers = tuple(self.df.columns)


    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        max_i = self.headers.index('Max Drift')
        avg_i = self.headers.index('Avg Drift')
        allow_i = self.headers.index('Allowable Drift')
        if index.isValid():
            value = self.df.iloc[row][col]
            allow_drift = float(self.df.iloc[row][allow_i])
            if role == Qt.DisplayRole:
                if col in (max_i, avg_i, allow_i):
                    return f"{value:.4f}"
                return value
            elif role == Qt.BackgroundColorRole:
                if col in (avg_i, max_i):
                    if float(value) > allow_drift:
                        return QColor(high)
                    else:
                        return QColor(low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)


class TorsionModel(ResultsModel):
    def __init__(self, data, headers):
        super(TorsionModel, self).__init__(data, headers)
        headers = [
            'Story',
            'Label',
            'OutputCase',
            'Max Drift',
            'Avg Drift',
            'Ratio',
        ]
        i_story = headers.index('Story')
        i_label = headers.index('Label')
        self.df = self.df[headers]
        self.headers = tuple(headers)
        self.col_function = (i_story, i_label)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        i_ratio = self.headers.index('Ratio')
        i_max = self.headers.index('Max Drift')
        i_avg = self.headers.index('Avg Drift')
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (i_max, i_avg, i_ratio):
                    return f"{value:.4f}"
                return str(value)
            elif role == Qt.BackgroundColorRole:
                value = float(self.df.iloc[row][i_ratio])
                # if col == i_ratio:
                    # value = float(value)
                if value <= 1.2:
                    return QColor(low)
                elif 1.2 < value < 1.4:
                    return QColor(intermediate)
                elif value > 1.4:
                    return QColor(high)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)


class StoryForcesModel(ResultsModel):
    def __init__(self, data, headers):
        super(StoryForcesModel, self).__init__(data, headers)
        self.df = self.df[[
            'Story',
            'OutputCase',
            'VX',
            'VY',
            'Vx %',
            'Vy %',
        ]]
        self.headers = tuple(self.df.columns)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        i_vx = self.headers.index('Vx %')
        i_vy = self.headers.index('Vy %')
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                return str(value)
            elif role == Qt.BackgroundColorRole:
                fx_Percentage = float(self.df.iloc[row][i_vx])
                fy_Percentage = float(self.df.iloc[row][i_vy])
                if max(fx_Percentage, fy_Percentage) >= .35:
                    return QColor(intermediate)
                else:
                    return QColor(low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)


class ColumnsRatioModel(ResultsModel):
    def __init__(self, data, headers):
        super(ColumnsRatioModel, self).__init__(data, headers)
        all_cols = list(self.df)
        self.df[all_cols] = self.df[all_cols].astype(str)
        self.headers = tuple(self.df.columns)
        # self.col_function = (0, 4)
    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                return str(value)
            elif role == Qt.BackgroundColorRole:
                ratio = float(self.df.iloc[row]['Ratio'])
                if ratio > 1:
                    return QColor(high)
                else:
                    return QColor(low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)


class BeamsRebarsModel(ResultsModel):
    def __init__(self, data, headers):
        super(BeamsRebarsModel, self).__init__(data, headers)
        all_cols = list(self.df)
        self.df[all_cols] = self.df[all_cols].astype(str)
        self.headers = tuple(self.df.columns)
        self.i_location = self.headers.index('location')
        self.i_ta1 = self.headers.index('Top Area1')
        self.i_ta2 = self.headers.index('Top Area2')
        self.i_ba1 = self.headers.index('Bot Area1')
        self.i_ba2 = self.headers.index('Bot Area2')
        self.i_v1 = self.headers.index('VRebar1')
        self.i_v2 = self.headers.index('VRebar2')
        # self.col_function = (0, 4)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (
                            self.i_ta1,
                            self.i_ta2,
                            self.i_ba1,
                            self.i_ba2,
                            ):
                    value = round(float(value), 1)
                if col in (
                            self.i_v1,
                            self.i_v2,
                            ):
                    value = round(float(value) * 100, 1)
                if col == self.i_location:
                    value = int(float(value))
                return str(value)
            elif role == Qt.BackgroundColorRole:
                if col in (self.i_ta1, self.i_ta2):
                    if float(self.df.iloc[row][self.i_ta2]) > float(self.df.iloc[row][self.i_ta1]) * 1.02:
                        return QColor(high)
                    else:
                        return QColor(low)
                if col in (self.i_ba1, self.i_ba2):
                    if float(self.df.iloc[row][self.i_ba2]) > float(self.df.iloc[row][self.i_ba1]) * 1.02:
                        return QColor(high)
                    else:
                        return QColor(low)
                if col in (self.i_v1, self.i_v2):
                    if float(self.df.iloc[row][self.i_v2]) > float(self.df.iloc[row][self.i_v1]) * 1.02:
                        return QColor(high)
                    else:
                        return QColor(low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)


class IrregularityOfMassModel(ResultsModel):
    def __init__(self, data, headers):
        super(IrregularityOfMassModel, self).__init__(data, headers)
        all_cols = list(self.df)
        self.df[all_cols] = self.df[all_cols].astype(str)
        self.headers = tuple(self.df.columns)
        self.i_mass_x = self.headers.index('Mass X')
        self.i_below = self.headers.index('1.5 * Below')
        self.i_above = self.headers.index('1.5 * Above')
        # self.col_function = (0, 4)
    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                return str(value)
            elif role == Qt.BackgroundColorRole:
                if col in (self.i_below, self.i_above):
                    if float(self.df.iloc[row][self.i_mass_x]) > \
                        float(self.df.iloc[row][col]):
                        return QColor(high)
                    else:
                        return QColor(low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)

class StoryStiffnessModel(ResultsModel):
    def __init__(self, data, headers):
        super(StoryStiffnessModel, self).__init__(data, headers)
        all_cols = list(self.df)
        self.df[all_cols] = self.df[all_cols].astype(str)
        self.headers = tuple(self.df.columns)
        self.i_kx = self.headers.index('Kx')
        self.i_ky = self.headers.index('Ky')
        self.i_kx_above = self.headers.index('Kx / kx+1')
        self.i_ky_above = self.headers.index('Ky / ky+1')
        self.i_kx_3above = self.headers.index('Kx / kx_3ave')
        self.i_ky_3above = self.headers.index('Ky / ky_3ave')
        # self.col_function = (0, 4)
    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (
                    self.i_kx_above,
                    self.i_kx_3above,
                    self.i_ky_above,
                    self.i_ky_3above,
                    ):
                    if value == '-':
                        return value
                    return f'{float(value):.3f}'
                elif col in (
                    self.i_kx,
                    self.i_ky,
                    ):
                    return f'{float(value):.0f}'
                return value
            elif role == Qt.BackgroundColorRole:
                if col in (self.i_kx_above, self.i_ky_above):
                    k = self.df.iloc[row][col]
                    return self.get_color(k, .6, .7)
                elif col in (self.i_kx_3above, self.i_ky_3above):
                    k = self.df.iloc[row][col]
                    return self.get_color(k, .7, .8)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)

    @staticmethod
    def get_color(k, a, b):
        if k == '-':
            return None
        k = float(k)
        if k < a:
            return QColor(high)
        elif k < b:
            return QColor(intermediate)
        else:
            return QColor(low)

class BeamsJModel(ResultsModel):
    def __init__(self, data, headers):
        super(BeamsJModel, self).__init__(data, headers)
        self.headers = tuple(self.df.columns)
        self.col_function = (2,)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        self.i_T = self.headers.index('T')
        self.i_Tcr = self.headers.index('phi_Tcr')
        self.i_j = self.headers.index('j')
        self.i_init_j = self.headers.index('init_j')
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (self.i_T, self.i_j, self.i_Tcr, self.i_init_j):
                    return f'{value:.2f}'
                return str(value)
            elif role == Qt.BackgroundColorRole:
                j = self.df.iloc[row][self.i_j]
                j_init = self.df.iloc[row][self.i_init_j]
                # if col == i_ratio:
                    # value = float(value)
                if j == j_init:
                    return QColor(low)
                else:
                    return QColor(intermediate)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)

class HighPressureColumnModel(ResultsModel):
    def __init__(self, data, headers):
        super(HighPressureColumnModel, self).__init__(data, headers)
        self.headers = tuple(self.df.columns)
        self.col_function = (3,)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        self.i_p = self.headers.index('P')
        self.i_t2 = self.headers.index('t2')
        self.i_t3 = self.headers.index('t3')
        self.i_fc = self.headers.index('fc')
        self.i_Agfc = self.headers.index('limit*Ag*fc')
        self.i_hp = self.headers.index('Result')
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (self.i_p, self.i_t2, self.i_t3, self.i_fc, self.i_Agfc):
                    return f'{value:.0f}'
                return str(value)
            elif role == Qt.BackgroundColorRole:
                if self.df.iloc[row][self.i_hp]:
                    return QColor(intermediate)
                else:
                    return QColor(low)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)

class Column100_30Model(ResultsModel):
    def __init__(self, data, headers):
        super(Column100_30Model, self).__init__(data, headers)
        self.headers = tuple(self.df.columns)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        self.i_p = self.headers.index('P')
        self.i_mmajor = self.headers.index('MMajor')
        self.i_mminor = self.headers.index('MMinor')
        self.i_ratio = self.headers.index('Ratio')
        self.i_result = self.headers.index('Result')
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (self.i_p, self.i_ratio):
                    return f'{value:.2f}'
                return str(value)
            elif role == Qt.BackgroundColorRole:
                if self.df.iloc[row][self.i_result]:
                    return QColor(low)
                else:
                    return QColor(high)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)

class JointShear(ResultsModel):
    def __init__(self, data, headers):
        super(JointShear, self).__init__(data, headers)
        self.headers = tuple(self.df.columns)
        self.i_maj = self.headers.index('JSMajRatio')
        self.i_min = self.headers.index('JSMinRatio')
        self.col_function = (2,)

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                if col in (self.i_maj, self.i_min):
                    return f'{value:.2f}'
                return str(value)
            elif role == Qt.BackgroundColorRole:
                if col in (self.i_maj, self.i_min):
                    if value <= 1:
                        return QColor(low)
                    else:
                        return QColor(high)
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)

class ExpandLoadSets(ResultsModel):
    def __init__(self, data, headers):
        super(ExpandLoadSets, self).__init__(data, headers)
        self.headers = tuple(self.df.columns)
        unique_names = self.df['UniqueName'].unique()
        self.i_uniquename = self.headers.index('UniqueName')
        
        self.colors = {}
        for name in unique_names:
            self.colors[name] = random_color()

    def data(self, index, role=Qt.DisplayRole):
        row = index.row()
        col = index.column()
        if index.isValid():
            value = self.df.iloc[row][col]
            if role == Qt.DisplayRole:
                return f'{value}'
            elif role == Qt.BackgroundColorRole:
                name = self.df.iloc[row][self.i_uniquename]
                return QColor.fromRgb(*self.colors[name])
            elif role == Qt.TextAlignmentRole:
                return int(Qt.AlignCenter | Qt.AlignVCenter)
            

class ResultWidget(QtWidgets.QDialog):
    # main widget for user interface
    def __init__(self, data, headers, model, function=None, parent=None):
        super(ResultWidget, self).__init__(parent)
        self.push_button_to_excel = QtWidgets.QPushButton()
        self.push_button_to_excel.setIcon(QIcon(str(civiltools_path / 'images' / 'xlsx.png')))
        self.push_button_to_word = QtWidgets.QPushButton()
        self.push_button_to_word.setIcon(QIcon(str(civiltools_path / 'images' / 'word.png')))
        label = QtWidgets.QLabel("Filter")
        self.lineEdit = QtWidgets.QLineEdit()
        label2 = QtWidgets.QLabel("By Column:")
        self.comboBox = QtWidgets.QComboBox()
        hbox = QtWidgets.QHBoxLayout()
        hbox.addWidget(label)
        hbox.addWidget(self.lineEdit)
        hbox.addWidget(label2)
        hbox.addWidget(self.comboBox)
        hbox.addWidget(self.push_button_to_word)
        hbox.addWidget(self.push_button_to_excel)
        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addLayout(hbox)
        self.result_table_view = QtWidgets.QTableView()
        self.result_table_view.setSortingEnabled(True)
        self.vbox.addWidget(self.result_table_view)
        self.setLayout(self.vbox)
        self.function = function
        self.data = data
        self.headers = headers
        self.model = model(self.data, self.headers)
        # self.result_table_view.setModel(self.model)
        self.proxy = QtCore.QSortFilterProxyModel(self)
        self.proxy.setSourceModel(self.model)
        self.result_table_view.setModel(self.proxy)
        self.comboBox.addItems(self.model.headers)
        self.lineEdit.textChanged.connect(self.on_lineEdit_textChanged)
        self.comboBox.currentIndexChanged.connect(self.on_comboBox_currentIndexChanged)
        # self.horizontalHeader = self.result_table_view.horizontalHeader()
        # self.horizontalHeader.sectionClicked.connect(self.on_view_horizontalHeader_sectionClicked)
        self.push_button_to_excel.clicked.connect(self.export_to_excel)
        self.push_button_to_word.clicked.connect(self.export_to_word)
        self.resize_columns()
        if self.function:
            self.result_table_view.clicked.connect(self.row_clicked)

    def row_clicked(self, index):
        source_index = self.proxy.mapToSource(index)
        row = source_index.row()
        # col = source_index.column()
        args = []
        for col in self.model.col_function:
            value = str(self.model.data(self.model.index(row, col)))
            args.append(value)
        self.function(*args) 

    def export_to_excel(self):
        filename, _ = QtWidgets.QFileDialog.getSaveFileName(self, 'export to excel',
                                                  '', "excel(*.xlsx)")
        if filename == '':
            return
        try:
            import jinja2
        except ModuleNotFoundError:
            import subprocess
            package = 'Jinja2'
            subprocess.check_call(['python', "-m", "pip", "install", package])
            css_alt_rows = 'background-color: powderblue; color: black;'
            css_indexes = 'background-color: steelblue; color: white;'
            import numpy as np
            (self.model.df.style.apply(lambda col: np.where(col.index % 2, css_alt_rows, None)) # alternating rows
                    .applymap_index(lambda _: css_indexes, axis=0) # row indexes (pandas 1.4.0+)
                    .applymap_index(lambda _: css_indexes, axis=1) # col indexes (pandas 1.4.0+)
            ).to_excel(filename, engine='openpyxl')
        else:
            with pd.ExcelWriter(filename) as writer:
                    self.model.df.to_excel(writer)
    
    def export_to_word(self,
                       ali='',
                       doc=None,
                       ):
        filename, _ = QtWidgets.QFileDialog.getSaveFileName(self, 'export to word',
                                                  '', "word(*.docx)")
        if filename == '':
            return
        try:
            from docx.shared import Inches
        except ImportError:
            package = 'python-docx'
            from freecad_funcs import install_package
            install_package(package)
    
        from docx import Document
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        # create a new document
        if doc is None:
            doc = Document()
        table_docx = doc.add_table(rows=self.model.rowCount()+1, cols=self.model.columnCount())
        table_docx.style = 'Table Grid'

        # write the column headers to the first row of the table
        for j in range(self.model.columnCount()):
            cell = table_docx.cell(0, j)
            cell.text = self.model.headerData(j, Qt.Horizontal, Qt.DisplayRole)
            # Set header text to bold
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.italic = True
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), "#244061"))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # write the data to the remaining rows of the table
        for row in range(self.model.rowCount()):
            for col in range(self.model.columnCount()):
                index = self.model.index(row, col)
                text = self.model.data(index)
                color = self.model.data(index, Qt.BackgroundColorRole)
                cell = table_docx.cell(row+1, col)
                cell.text = text
                if color:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color.name()))
                    cell._tc.get_or_add_tcPr().append(shading_elm)


        # save the document
        doc.save(filename)
        # self.model.df.to_excel()

    def resize_columns(self):
        self.result_table_view.resizeColumnsToContents()
        self.adjustSize()
        # width = 0
        # for col in range(len(self.model.df.columns)):
        #     width += self.result_table_view.columnWidth(col)
        # self.result_table_view.setFixedWidth(width)


    # @QtCore.Slot(int)
    # def on_view_horizontalHeader_sectionClicked(self, logicalIndex):
    #     self.logicalIndex   = logicalIndex
    #     self.menuValues     = QtWidgets.QMenu(self)
    #     self.signalMapper   = QtCore.QSignalMapper(self)  

    #     self.comboBox.blockSignals(True)
    #     self.comboBox.setCurrentIndex(self.logicalIndex)
    #     self.comboBox.blockSignals(True)


    #     valuesUnique = list(self.model.df.iloc[:, self.logicalIndex].unique())

    #     actionAll = QtWidgets.QAction("All", self)
    #     actionAll.triggered.connect(self.on_actionAll_triggered)
    #     self.menuValues.addAction(actionAll)
    #     self.menuValues.addSeparator()

    #     for actionNumber, actionName in enumerate(sorted(list(set(valuesUnique)))):              
    #         action = QtWidgets.QAction(actionName, self)
    #         self.signalMapper.setMapping(action, actionNumber)  
    #         action.triggered.connect(self.signalMapper.map)  
    #         self.menuValues.addAction(action)

    #     self.signalMapper.mapped.connect(self.on_signalMapper_mapped)  

    #     headerPos = self.result_table_view.mapToGlobal(self.horizontalHeader.pos())        

    #     posY = headerPos.y() + self.horizontalHeader.height()
    #     posX = headerPos.x() + self.horizontalHeader.sectionPosition(self.logicalIndex)

    #     self.menuValues.exec_(QtCore.QPoint(posX, posY))

    @QtCore.Slot()
    def on_actionAll_triggered(self):
        filterColumn = self.logicalIndex
        filterString = QtCore.QRegExp(  "",
                                        QtCore.Qt.CaseInsensitive,
                                        QtCore.QRegExp.RegExp
                                        )

        self.proxy.setFilterRegExp(filterString)
        self.proxy.setFilterKeyColumn(filterColumn)

    @QtCore.Slot(int)
    def on_signalMapper_mapped(self, i):
        stringAction = self.signalMapper.mapping(i).text()
        filterColumn = self.logicalIndex
        filterString = QtCore.QRegExp(  stringAction,
                                        QtCore.Qt.CaseSensitive,
                                        QtCore.QRegExp.FixedString
                                        )

        self.proxy.setFilterRegExp(filterString)
        self.proxy.setFilterKeyColumn(filterColumn)

    @QtCore.Slot(str)
    def on_lineEdit_textChanged(self, text):
        search = QtCore.QRegExp(    text,
                                    QtCore.Qt.CaseInsensitive,
                                    QtCore.QRegExp.RegExp
                                    )

        self.proxy.setFilterRegExp(search)

    @QtCore.Slot(int)
    def on_comboBox_currentIndexChanged(self, index):
        self.proxy.setFilterKeyColumn(index)


class ExpandedLoadSetsResults(ResultWidget):
    def __init__(self, data, headers, model, function=None, parent=None):
        super(ExpandedLoadSetsResults, self).__init__(data, headers, model, function, parent)
        self.cancel_pushbutton = QtWidgets.QPushButton()
        self.cancel_pushbutton.setIcon(QIcon(str(civiltools_path / 'images' / 'cancel.svg')))
        self.cancel_pushbutton.setText('&Cancel')
        self.apply_pushbutton = QtWidgets.QPushButton()
        self.apply_pushbutton.setIcon(QIcon(str(civiltools_path / 'images' / 'etabs.png')))
        self.apply_pushbutton.setText('&Apply')

        hbox = QtWidgets.QHBoxLayout()
        hbox.addWidget(self.cancel_pushbutton)
        hbox.addWidget(self.apply_pushbutton)
        self.vbox.addLayout(hbox)

def show_results(data, headers, model, function=None):
    win = ResultWidget(data, headers, model, function)
    # Gui.Control.showDialog(win)
    mdi = get_mdiarea()
    if not mdi:
        return None
    sub = mdi.addSubWindow(win)
    sub.show()

def get_mdiarea():
    """ Return FreeCAD MdiArea. """
    import FreeCADGui as Gui
    import PySide2
    mw = Gui.getMainWindow()
    if not mw:
        return None
    childs = mw.children()
    for c in childs:
        if isinstance(c, PySide2.QtWidgets.QMdiArea):
            return c
    return None

def random_color():
    h,s,l = random.random(), 0.5 + random.random()/2.0, 0.4 + random.random()/5.0
    return [int(256*i) for i in colorsys.hls_to_rgb(h,l,s)]


# class EtabsModel:
#     # my ETABS API class to open and manipulate etabs model
#     def __init__(self, modelpath, etabspath="c:/Program Files/Computers and Structures/ETABS 19/ETABS.exe", existinstance=False, specprogpath=False):
#         # set the following flag to True to attach to an existing instance of the program
#         # otherwise a new instance of the program will be started
#         self.AttachToInstance = existinstance

#         # set the following flag to True to manually specify the path to ETABS.exe
#         # this allows for a connection to a version of ETABS other than the latest installation
#         # otherwise the latest installed version of ETABS will be launched
#         self.SpecifyPath = specprogpath

#         # if the above flag is set to True, specify the path to ETABS below
#         self.ProgramPath = etabspath

#         # full path to the model
#         # set it to the desired path of your model
#         self.FullPath = modelpath
#         [self.modelPath, self.modelName] = os.path.split(self.FullPath)
#         if not os.path.exists(self.modelPath):
#             try:
#                 os.makedirs(self.modelPath)
#             except OSError:
#                 pass

#         if self.AttachToInstance:
#             # attach to a running instance of ETABS
#             try:
#                 # get the active ETABS object
#                 self.myETABSObject = comtypes.client.GetActiveObject("CSI.ETABS.API.ETABSObject")
#                 self.success = True
#             except (OSError, comtypes.COMError):
#                 print("No running instance of the program found or failed to attach.")
#                 self.success = False
#                 sys.exit(-1)

#         else:
#             # create API helper object
#             self.helper = comtypes.client.CreateObject('ETABSv1.Helper')
#             self.helper = self.helper.QueryInterface(comtypes.gen.ETABSv1.cHelper)
#             if self.SpecifyPath:
#                 try:
#                     # 'create an instance of the ETABS object from the specified path
#                     self.myETABSObject = self.helper.CreateObject(self.ProgramPath)
#                 except (OSError, comtypes.COMError):
#                     print("Cannot start a new instance of the program from " + self.ProgramPath)
#                     sys.exit(-1)
#             else:

#                 try:
#                     # create an instance of the ETABS object from the latest installed ETABS
#                     self.myETABSObject = self.helper.CreateObjectProgID("CSI.ETABS.API.ETABSObject")
#                 except (OSError, comtypes.COMError):
#                     print("Cannot start a new instance of the program.")
#                     sys.exit(-1)

#             # start ETABS application
#             self.myETABSObject.ApplicationStart()

#         # create SapModel object
#         self.SapModel = self.myETABSObject.SapModel

#         # initialize model
#         self.SapModel.InitializeNewModel()

#         # create new blank model
#         # ret = self.SapModel.File.NewBlank()

#         # open existing model
#         ret = self.SapModel.File.OpenFile(self.FullPath)

#         """
#         # save model
#         print(self.FullPath)
#         ret = self.SapModel.File.Save(self.FullPath)
#         print(ret)
#         print("ETABS mod - model saved")
#         """

#         # run model (this will create the analysis model)
#         ret = self.SapModel.Analyze.RunAnalysis()

#         # get all load combination names
#         self.NumberCombo = 0
#         self.ComboNames = []
#         [self.NumberCombo, self.ComboNames, ret] = self.SapModel.RespCombo.GetNameList(self.NumberCombo, self.ComboNames)

#         # isolate drift combos by searching for "drift" in combo name
#         self.DriftCombos = []
#         for combo in self.ComboNames:
#             lowerCombo = combo.lower()
#             # skip combinations without drift in name
#             if "drift" not in lowerCombo:
#                 continue
#             self.DriftCombos.append(combo)

#         self.StoryDrifts = []
#         self.JointDisplacements = []
#         pd.set_option("max_columns", 8)
#         # pd.set_option("precision", 4)

#     def story_drift_results(self, dlimit=0.01):
#         # returns dataframe drift results for all drift load combinations
#         self.StoryDrifts = []
#         for dcombo in self.DriftCombos:
#             # deselect all combos and cases, then only display results for combo passed
#             ret = self.SapModel.Results.Setup.DeselectAllCasesAndCombosForOutput()
#             ret = self.SapModel.Results.Setup.SetComboSelectedForOutput(dcombo)

#             # initialize drift results
#             NumberResults = 0
#             Stories = []
#             LoadCases = []
#             StepTypes = []
#             StepNums = []
#             Directions = []
#             Drifts = []
#             Labels = []
#             Xs = []
#             Ys = []
#             Zs = []

#             [NumberResults, Stories, LoadCases, StepTypes, StepNums, Directions, Drifts, Labels, Xs, Ys, Zs, ret] = \
#                 self.SapModel.Results.StoryDrifts(NumberResults, Stories, LoadCases, StepTypes, StepNums, Directions,
#                                                   Drifts, Labels, Xs, Ys, Zs)
#             # append all drift results to storydrifts list
#             for i in range(0, NumberResults):
#                 self.StoryDrifts.append((Stories[i], LoadCases[i], Directions[i], Drifts[i], Drifts[i] / dlimit))

#         # set up pandas data frame and sort by drift column
#         labels = ['Story', 'Combo', 'Direction', 'Drift', 'DCR(Drift/Limit)']
#         df = pd.DataFrame.from_records(self.StoryDrifts, columns=labels)
#         dfSort = df.sort_values(by=['Drift'], ascending=False)
#         dfSort.Drift = dfSort.Drift.round(4)
#         dfSort['DCR(Drift/Limit)'] = dfSort['DCR(Drift/Limit)'].round(2)
#         return dfSort

#     def story_torsion_check(self):
#         # returns dataframe of torsion results for drift combinations
#         self.JointDisplacements = []
#         for dcombo in self.DriftCombos:
#             # deselect all combos and cases, then only display results for combo passed
#             ret = self.SapModel.Results.Setup.DeselectAllCasesAndCombosForOutput()
#             ret = self.SapModel.Results.Setup.SetComboSelectedForOutput(dcombo)

#             # initialize joint drift results
#             NumberResults = 0
#             Stories = []
#             LoadCases = []
#             Label  = ''
#             Names = ''
#             StepType = []
#             StepNum = []
#             # Directions = []
#             DispX = []
#             DispY = []
#             DriftX = []
#             DriftY = []

#             [NumberResults, Stories, Label, Names, LoadCases, StepType, StepNum, DispX, DispY, DriftX, DriftY, ret] = \
#                 self.SapModel.Results.JointDrifts(NumberResults, Stories, Label, Names, LoadCases, StepType, StepNum,
#                                                   DispX, DispY, DriftX, DriftY)

#             # append all displacement results to jointdrift list
#             for i in range(0, NumberResults):
#                 self.JointDisplacements.append((Label[i], Stories[i], LoadCases[i], DispX[i], DispY[i]))

#         # set up pandas data frame and sort by drift column
#         jlabels = ['label', 'Story', 'Combo', 'DispX', 'DispY']
#         jdf = pd.DataFrame.from_records(self.JointDisplacements, columns=jlabels)
#         story_names = jdf.Story.unique()
#         # print("story names = " + str(story_names))

#         # set up data frame for torsion displacement results
#         tlabels = ['Story', 'Load Combo', 'Direction', 'Max Displ', 'Avg Displ', 'Ratio']
#         tdf = pd.DataFrame(columns=tlabels)

#         # calculate and append to df the max, avg, and ratio of story displacements in each dir.
#         for dcombo in self.DriftCombos:
#             for story in story_names:
#                 temp_df = jdf[(jdf['Story'] == story) & (jdf['Combo'] == dcombo)]
#                 # assume direction is X
#                 direction = 'X'
#                 averaged = abs(temp_df['DispX'].mean())
#                 maximumd = temp_df['DispX'].abs().max()
#                 averagey = abs(temp_df['DispY'].mean())

#                 # change direction to Y if avg y-dir displacement is higher
#                 if averagey > averaged:
#                     averaged = averagey
#                     maximumd = temp_df['DispY'].abs().max()
#                     direction = 'Y'

#                 ratiod = maximumd / averaged

#                 # append info to torsion dataframe
#                 temp_df2 = pd.DataFrame([[story, dcombo, direction, maximumd, averaged, ratiod]], columns=tlabels)
#                 tdf = tdf.append(temp_df2, ignore_index=True)

#         tdfSort = tdf.sort_values(by=['Ratio'], ascending=False)
#         tdfSort.Ratio = tdfSort.Ratio.round(3)
#         tdfSort['Max Displ'] = tdfSort['Max Displ'].round(3)
#         tdfSort['Avg Displ'] = tdfSort['Avg Displ'].round(3)

#         return tdfSort

#     def model_close(self):
#         # close the program
#         ret = self.myETABSObject.ApplicationExit(False)
#         self.SapModel = None
#         self.myETABSObject = None
#         return 'model closed'

# """
# # define drift limit for DCR calc
# DriftLimit = 0.01
# FilePath = "C:\\Users\\Andrew-V.Young\\Desktop\\ETABS API TEST\\ETABS\\TestModel.EDB"

# testModel = EtabsModel(FilePath)
# drifts = testModel.story_drift_results(DriftLimit)
# torsion = testModel.story_torsion_check()
# print(drifts.head(10))
# print("\n\n")
# print(torsion)

# ret = testModel.model_close()
# print(ret)